"""Genera el manual de usuario de Terralix ERP como .docx"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

OUT = os.path.join(os.path.dirname(__file__), "Manual_Terralix_ERP.docx")

doc = Document()

# --- Estilos globales ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.color.rgb = RGBColor(0x0F, 0x66, 0x45)  # verde Terralix
    hs.font.name = "Calibri"

# --- Funciones auxiliares ---
def add_table_row(table, cells, bold=False, bg=None):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if bold:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if bg:
            shading = cell._element.get_or_add_tcPr()
            sh = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): bg,
            })
            shading.append(sh)


def make_table(headers, rows_data, col_widths=None):
    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Light Grid Accent 1"
    add_table_row(table, headers, bold=True, bg="0F6645")
    for row_data in rows_data:
        add_table_row(table, row_data)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)
    return table


def add_note(text, color_hex="FFF9C4"):
    """Agrega un parrafo con fondo de nota/aviso."""
    p = doc.add_paragraph()
    run = p.add_run(f"\u2139\ufe0f  {text}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.name = "Calibri"
    return p


def add_warning(text):
    """Agrega un parrafo de advertencia."""
    p = doc.add_paragraph()
    run_lbl = p.add_run("\u26a0\ufe0f  Importante:  ")
    run_lbl.bold = True
    run_lbl.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    run_lbl.font.name = "Calibri"
    run_lbl.font.size = Pt(10)
    run_txt = p.add_run(text)
    run_txt.font.size = Pt(10)
    run_txt.font.name = "Calibri"
    return p


# =====================================================================
# PORTADA
# =====================================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.space_before = Pt(120)
run = p.add_run("TERRALIX ERP")
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x0F, 0x66, 0x45)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Manual de Usuario")
run2.font.size = Pt(24)
run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.space_before = Pt(40)
run3 = p3.add_run("Versi\u00f3n 1.3.0\nMarzo 2026")
run3.font.size = Pt(14)
run3.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p4.space_before = Pt(20)
run4 = p4.add_run("Agr\u00edcola Las Tipuanas SPA")
run4.font.size = Pt(12)
run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run4.italic = True

doc.add_page_break()

# =====================================================================
# ÍNDICE
# =====================================================================
doc.add_heading("\u00cdndice", level=1)
toc_items = [
    "1. Introducci\u00f3n",
    "2. Requisitos del sistema",
    "3. Instalaci\u00f3n y primera ejecuci\u00f3n",
    "   3.1 Instalaci\u00f3n",
    "   3.2 Primera ejecuci\u00f3n",
    "4. Inicio de sesi\u00f3n",
    "   4.1 Iniciar sesi\u00f3n",
    "   4.2 Recordar credenciales",
    "   4.3 Recuperar contrase\u00f1a",
    "5. Pantalla principal",
    "   5.1 Pesta\u00f1a Actualizar",
    "   5.2 Pesta\u00f1a Inventario",
    "   5.3 Pesta\u00f1a Aplicaciones",
    "   5.4 Barra de men\u00fa superior",
    "6. Actualizar Base de Datos (pipeline completo \u2014 6 etapas)",
    "   Etapa 1: Descarga desde SII",
    "   Etapa 2: Lectura IA de PDFs",
    "   Etapa 3: Categorizaci\u00f3n contable",
    "   Etapa 4: Backfill de c\u00f3digos (INSUMOS_AGR\u00cdCOLAS)",
    "   Etapa 5: Backfill de unidades",
    "   Etapa 6: Sincronizaci\u00f3n autom\u00e1tica de inventario",
    "7. Exportar Excel para revisi\u00f3n",
    "8. Correcci\u00f3n manual en Excel",
    "9. Importar Excel y reentrenar modelo",
    "10. M\u00f3dulo de Inventario",
    "    10.1 Vista de stock actual",
    "    10.2 Registrar uso manual",
    "    10.3 Edici\u00f3n directa de celdas",
    "    10.4 Sincronizar desde DTE",
    "11. M\u00f3dulo de Aplicaciones de campo",
    "    11.1 Crear una nueva aplicaci\u00f3n",
    "    11.2 Agregar productos qu\u00edmicos",
    "    11.3 Calendario mensual",
    "    11.4 Gestionar aplicaciones existentes",
    "12. Chequeo semanal autom\u00e1tico en segundo plano",
    "13. Configurar rutas",
    "14. Conexi\u00f3n con Power BI / Excel externo",
    "15. Estructura de la base de datos",
    "16. Cat\u00e1logo de costos",
    "17. Preguntas frecuentes",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =====================================================================
# 1. INTRODUCCIÓN
# =====================================================================
doc.add_heading("1. Introducci\u00f3n", level=1)
doc.add_paragraph(
    "Terralix ERP es una aplicaci\u00f3n de escritorio desarrollada en Python "
    "para la gesti\u00f3n automatizada de Documentos Tributarios Electr\u00f3nicos (DTE) "
    "recibidos por Agr\u00edcola Las Tipuanas SPA. Combina descarga autom\u00e1tica desde "
    "el SII, lectura inteligente de PDFs con GPT-4o, clasificaci\u00f3n contable con "
    "Machine Learning local y un m\u00f3dulo completo de inventario de insumos agr\u00edcolas."
)
doc.add_paragraph(
    "La aplicaci\u00f3n cubre el ciclo completo de gesti\u00f3n de facturas y boletas:"
)
items = [
    "Descarga autom\u00e1tica desde el Servicio de Impuestos Internos (SII)",
    "Lectura inteligente de PDFs mediante visi\u00f3n artificial (GPT-4o)",
    "Categorizaci\u00f3n contable autom\u00e1tica con modelo de Machine Learning local",
    "Completado autom\u00e1tico de c\u00f3digos de insumos y unidades de medida",
    "Sincronizaci\u00f3n de inventario te\u00f3rico desde los DTE comprados",
    "Exportaci\u00f3n a Excel para revisi\u00f3n manual y correcci\u00f3n",
    "Aprendizaje continuo: las correcciones manuales reentrenan el modelo",
    "M\u00f3dulo de aplicaciones de campo con calendario mensual",
    "Conexi\u00f3n directa con Power BI y Excel para an\u00e1lisis y reportes",
    "Chequeo semanal autom\u00e1tico en segundo plano",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_paragraph(
    "El objetivo es que con un solo clic se actualice toda la base de datos, "
    "permiti\u00e9ndote enfocarte en el an\u00e1lisis y la gesti\u00f3n, no en la carga manual de datos."
)

# =====================================================================
# 2. REQUISITOS
# =====================================================================
doc.add_heading("2. Requisitos del sistema", level=1)
make_table(
    ["Componente", "Requisito"],
    [
        ["Sistema operativo", "Windows 10 / 11 (64 bits)"],
        ["Python", "3.11 o superior (incluido en el instalador)"],
        ["RAM", "4 GB m\u00ednimo, 8 GB recomendado"],
        ["Disco", "500 MB para la app + espacio para PDFs"],
        ["Internet", "Requerido para descarga SII, lectura IA e inicio de sesi\u00f3n"],
        ["Navegador", "Chromium (instalado autom\u00e1ticamente por Playwright)"],
        ["Excel", "Microsoft Excel o LibreOffice Calc (para revisi\u00f3n)"],
        ["Cuenta Terralix", "Email y contrase\u00f1a proporcionados por el administrador"],
    ],
    col_widths=[5, 12],
)

# =====================================================================
# 3. INSTALACIÓN
# =====================================================================
doc.add_heading("3. Instalaci\u00f3n y primera ejecuci\u00f3n", level=1)

doc.add_heading("3.1 Instalaci\u00f3n", level=2)
doc.add_paragraph(
    "Descarga el instalador TERRALIX_Setup.exe y ejec\u00fatalo con doble clic. "
    "El asistente instalar\u00e1 la aplicaci\u00f3n en C:\\Program Files\\Terralix ERP\\ "
    "y crear\u00e1 un acceso directo en el escritorio y el men\u00fa Inicio."
)
add_note(
    "Los datos del usuario (base de datos, PDFs, config.env, logs) se guardan en "
    "%APPDATA%\\Terralix ERP\\ y no se eliminan al desinstalar la aplicaci\u00f3n."
)

doc.add_heading("3.2 Primera ejecuci\u00f3n", level=2)
doc.add_paragraph(
    "En la primera ejecuci\u00f3n, la aplicaci\u00f3n mostrar\u00e1 un asistente de configuraci\u00f3n "
    "inicial que te pedir\u00e1 definir dos rutas esenciales:"
)
make_table(
    ["Ruta", "Descripci\u00f3n", "Ejemplo"],
    [
        ["Carpeta de PDFs", "Donde se guardar\u00e1n los PDF descargados del SII", "C:\\Users\\...\\DTE"],
        ["Carpeta de base de datos", "Donde se crear\u00e1 DteRecibidos_db.db", "C:\\Users\\...\\Base"],
    ],
    col_widths=[4.5, 8, 4.5],
)
doc.add_paragraph(
    "Estas rutas se guardan en config.env y pueden cambiarse posteriormente "
    "desde el men\u00fa Opciones > Configurar Rutas (PDF / Base de Datos)."
)

# =====================================================================
# 4. INICIO DE SESIÓN
# =====================================================================
doc.add_heading("4. Inicio de sesi\u00f3n", level=1)
doc.add_paragraph(
    "Al abrir Terralix ERP aparece la ventana de ingreso. El sistema utiliza "
    "autenticaci\u00f3n segura a trav\u00e9s de Supabase con email y contrase\u00f1a. "
    "Solo los usuarios autorizados por el administrador pueden acceder."
)

doc.add_heading("4.1 Iniciar sesi\u00f3n", level=2)
steps = [
    "Escribe tu email en el campo correspondiente.",
    "Escribe tu contrase\u00f1a.",
    "Presiona el bot\u00f3n Continuar o la tecla Enter.",
    "Mientras se verifica, aparecer\u00e1 el mensaje \u2018Conectando...\u2019",
    "Si las credenciales son incorrectas, se muestra un mensaje de error en rojo.",
]
for step in steps:
    doc.add_paragraph(step, style="List Bullet")

make_table(
    ["Atajo", "Acci\u00f3n"],
    [
        ["Enter / KP_Enter", "Iniciar sesi\u00f3n (equivale a presionar Continuar)"],
        ["Escape", "Salir de la aplicaci\u00f3n"],
    ],
    col_widths=[5, 12],
)

doc.add_heading("4.2 Recordar credenciales", level=2)
doc.add_paragraph(
    "Si marcas la casilla \u2018Recordar credenciales\u2019, la aplicaci\u00f3n guardar\u00e1 tu email y "
    "contrase\u00f1a de forma segura en el Administrador de Credenciales de Windows "
    "(Windows Credential Manager). La pr\u00f3xima vez que abras la app, los campos "
    "se rellenar\u00e1n autom\u00e1ticamente."
)

doc.add_heading("4.3 Recuperar contrase\u00f1a", level=2)
doc.add_paragraph(
    "Si olvidaste tu contrase\u00f1a, sigue estos pasos:"
)
steps = [
    "Escribe tu email en el campo correspondiente.",
    "Haz clic en el enlace azul \u2018\u00bfOlvidaste tu contrase\u00f1a?\u2019",
    "Revisa tu bandeja de entrada (y la carpeta spam).",
    "Haz clic en el enlace del correo de recuperaci\u00f3n.",
    "Se abrir\u00e1 autom\u00e1ticamente una ventana en la app para escribir tu nueva contrase\u00f1a.",
    "Ingresa y confirma la nueva contrase\u00f1a (m\u00ednimo 6 caracteres).",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Paso {i}: ")
    run.bold = True
    p.add_run(step)

# =====================================================================
# 5. PANTALLA PRINCIPAL
# =====================================================================
doc.add_heading("5. Pantalla principal", level=1)
doc.add_paragraph(
    "Tras iniciar sesi\u00f3n, la ventana de login se transforma en la aplicaci\u00f3n principal. "
    "La pantalla es redimensionable y se adapta autom\u00e1ticamente al tama\u00f1o de la ventana. "
    "La interfaz utiliza tres pesta\u00f1as en la parte superior."
)

doc.add_heading("5.1 Pesta\u00f1a Actualizar", level=2)
doc.add_paragraph(
    "Es la pesta\u00f1a principal. Contiene a la izquierda el logo de Terralix y el bot\u00f3n "
    "Actualizar que ejecuta el pipeline completo de 6 etapas. A la derecha hay una "
    "consola de texto que muestra el progreso en tiempo real y dos barras de progreso: "
    "una para las descargas y otra para la lectura de PDFs."
)

doc.add_heading("5.2 Pesta\u00f1a Inventario", level=2)
doc.add_paragraph(
    "M\u00f3dulo de gesti\u00f3n de inventario de insumos agr\u00edcolas. Muestra el stock actual "
    "de cada producto, permite registrar usos manuales y editar informaci\u00f3n del cat\u00e1logo. "
    "Ver secci\u00f3n 10 para el detalle completo."
)

doc.add_heading("5.3 Pesta\u00f1a Aplicaciones", level=2)
doc.add_paragraph(
    "M\u00f3dulo de planificaci\u00f3n de aplicaciones de campo (tratamientos fitosanitarios, "
    "fertilizaciones, etc.). Incluye un calendario mensual y la posibilidad de registrar "
    "los productos qu\u00edmicos utilizados en cada aplicaci\u00f3n. Ver secci\u00f3n 11 para el detalle."
)

doc.add_heading("5.4 Barra de men\u00fa superior", level=2)
make_table(
    ["Men\u00fa", "Opci\u00f3n", "Funci\u00f3n"],
    [
        ["Excel", "Exportar Excel", "Genera DteRecibidos_revision.xlsx con todos los datos para revisi\u00f3n"],
        ["Excel", "Importar Excel", "Lee correcciones del Excel, actualiza la DB y reentrena el modelo ML"],
        ["Inventario", "Refrescar Inventario", "Sincroniza el inventario desde los DTE (equivale al bot\u00f3n Refrescar del tab)"],
        ["Aplicaciones", "Refrescar Aplicaciones", "Recarga el calendario y la lista de aplicaciones de campo"],
        ["Opciones", "Configurar Rutas (PDF / Base de Datos)", "Permite cambiar la carpeta de PDFs y la ruta de la base de datos"],
        ["Ayuda", "Manual de usuario", "Abre este manual en formato PDF"],
    ],
    col_widths=[3, 5.5, 8.5],
)

# =====================================================================
# 6. PIPELINE (6 ETAPAS)
# =====================================================================
doc.add_heading("6. Actualizar Base de Datos (pipeline completo \u2014 6 etapas)", level=1)
doc.add_paragraph(
    "El bot\u00f3n Actualizar ejecuta las 6 etapas del pipeline de forma secuencial. "
    "Cada etapa se muestra en la consola con el formato [N/6]. "
    "No cierres la ventana hasta que aparezca el mensaje "
    "\u2018Flujo completo terminado correctamente.\u2019"
)
add_warning(
    "Si otro proceso del pipeline ya est\u00e1 en ejecuci\u00f3n (por ejemplo el chequeo "
    "semanal autom\u00e1tico), el bot\u00f3n Actualizar mostrar\u00e1 un aviso y no iniciar\u00e1 "
    "hasta que el proceso anterior termine."
)

doc.add_heading("Etapa 1: Descarga desde SII", level=2)
doc.add_paragraph(
    "Se conecta autom\u00e1ticamente al Servicio de Impuestos Internos usando las "
    "credenciales configuradas en config.env (RUT y CLAVE). Un navegador Chromium "
    "invisible navega por el portal del SII, p\u00e1gina por p\u00e1gina, y descarga los PDF "
    "de facturas y boletas recibidas que a\u00fan no est\u00e9n en la carpeta local."
)
items = [
    "La barra Descargas muestra el progreso (N / total).",
    "Si ya existe el PDF, se omite para evitar duplicados.",
    "Si la descarga falla, se contin\u00faa con los archivos disponibles.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Etapa 2: Lectura IA de PDFs", level=2)
doc.add_paragraph(
    "Cada PDF nuevo se convierte a im\u00e1genes y se env\u00eda a GPT-4o (visi\u00f3n) que extrae "
    "la informaci\u00f3n estructurada: emisor, folio, fecha, montos y el detalle l\u00ednea por "
    "l\u00ednea (descripci\u00f3n, cantidad, precio unitario, monto). Los datos se guardan en "
    "las tablas documentos y detalle."
)
items = [
    "La barra Lectura PDF muestra el porcentaje completado.",
    "Si un PDF falla, se reintenta hasta 4 veces con espera entre reintentos.",
    "Los documentos ya existentes en la DB se omiten autom\u00e1ticamente.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("Etapa 3: Categorizaci\u00f3n contable", level=2)
doc.add_paragraph(
    "Cada l\u00ednea de detalle se clasifica autom\u00e1ticamente usando un sistema de 3 capas:"
)
items_clas = [
    ("Capa 1 \u2013 Reglas locales",
     "Patrones de proveedor y descripci\u00f3n (ej.: si el proveedor es una autopista, "
     "la categor\u00eda es MANTENCI\u00d3N > VEH\u00cdCULOS > PEAJES). Confianza: 90-95%."),
    ("Capa 2 \u2013 Mantenedor de proveedores",
     "Si el proveedor ya fue clasificado antes con alta confianza, "
     "se reutiliza esa categor\u00eda. Confianza: 92%."),
    ("Capa 3 \u2013 Modelo ML local",
     "Un clasificador TF-IDF + LinearSVC entrenado con los datos hist\u00f3ricos. "
     "Usa la descripci\u00f3n, proveedor, giro y temporada agr\u00edcola para predecir la categor\u00eda. "
     "Funciona sin internet y sin costo de API."),
]
for title, desc in items_clas:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.bold = True
    p.add_run(desc)

doc.add_heading("Etapa 4: Backfill de c\u00f3digos (INSUMOS_AGR\u00cdCOLAS)", level=2)
doc.add_paragraph(
    "Despu\u00e9s de clasificar, vuelve a leer los PDFs de la categor\u00eda INSUMOS_AGR\u00cdCOLAS "
    "para extraer el c\u00f3digo del producto (ej.: c\u00f3digo de agroqu\u00edmico o fertilizante). "
    "Esta informaci\u00f3n es cr\u00edtica para el m\u00f3dulo de inventario. "
    "Solo procesa los documentos que a\u00fan no tienen c\u00f3digo asignado."
)

doc.add_heading("Etapa 5: Backfill de unidades", level=2)
doc.add_paragraph(
    "Completa la columna unidad (litros, kg, saco, etc.) para los INSUMOS_AGR\u00cdCOLAS "
    "que no la tienen, usando los datos de debug JSON guardados localmente durante la "
    "lectura IA. No requiere llamadas a la API de OpenAI."
)

doc.add_heading("Etapa 6: Sincronizaci\u00f3n autom\u00e1tica de inventario", level=2)
doc.add_paragraph(
    "Sincroniza el inventario te\u00f3rico de insumos a partir de los datos de detalle. "
    "Solo considera documentos tipo Factura y productos con c\u00f3digo asignado "
    "(ignora c\u00f3digo=\u2018-\u2019). Actualiza el cat\u00e1logo de insumos y los movimientos de "
    "entrada en el m\u00f3dulo de Inventario."
)

# =====================================================================
# 7. EXPORTAR EXCEL
# =====================================================================
doc.add_heading("7. Exportar Excel para revisi\u00f3n", level=1)
doc.add_paragraph(
    "Ve al men\u00fa Excel > Exportar Excel. La aplicaci\u00f3n genera el archivo "
    "DteRecibidos_revision.xlsx en la misma carpeta de la base de datos."
)

doc.add_heading("Estructura del Excel", level=2)
make_table(
    ["Hoja", "Contenido", "Editable"],
    [
        ["detalle", "Todas las l\u00edneas de detalle con su clasificaci\u00f3n", "S\u00ed (columnas de categor\u00eda)"],
        ["cat\u00e1logo", "Las combinaciones v\u00e1lidas de categor\u00eda/subcategor\u00eda/tipo_gasto", "No (solo referencia)"],
        ["documentos", "Datos de los documentos: emisor, fecha, montos", "No (solo referencia)"],
        ["stock", "Stock actual de insumos por c\u00f3digo", "No (solo referencia)"],
        ["entradas", "Movimientos de entrada de insumos desde facturas", "No (solo referencia)"],
    ],
    col_widths=[3, 10, 4],
)

doc.add_heading("Codificaci\u00f3n de colores", level=2)
make_table(
    ["Color", "Significado"],
    [
        ["Amarillo", "Filas marcadas como \u2018needs_review\u2019: el modelo no est\u00e1 seguro de la clasificaci\u00f3n"],
        ["Rojo claro", "Filas SIN_CLASIFICAR: no se pudo determinar la categor\u00eda autom\u00e1ticamente"],
        ["Verde", "Filas corregidas manualmente (origen=MANUAL): ya validadas por el usuario"],
        ["Gris (texto)", "Columnas de solo lectura (no editables)"],
    ],
    col_widths=[4, 13],
)
doc.add_paragraph(
    "La columna \u2018categor\u00eda\u2019 tiene un desplegable con las categor\u00edas v\u00e1lidas del cat\u00e1logo."
)

# =====================================================================
# 8. CORRECCIÓN MANUAL
# =====================================================================
doc.add_heading("8. Correcci\u00f3n manual en Excel", level=1)
doc.add_paragraph("Para corregir clasificaciones err\u00f3neas:")

steps = [
    "Abre el Excel generado (DteRecibidos_revision.xlsx)",
    "Usa los filtros para mostrar solo filas con needs_review=1 o categor\u00eda=\u2018SIN_CLASIFICAR\u2019",
    "Consulta la hoja \u2018cat\u00e1logo\u2019 para ver las combinaciones v\u00e1lidas",
    "Modifica las columnas \u2018categor\u00eda\u2019, \u2018subcategor\u00eda\u2019 y \u2018tipo_gasto\u2019 seg\u00fan corresponda",
    "Guarda el archivo Excel (Ctrl+S)",
    "Vuelve a Terralix y ve a Excel > Importar Excel",
]
for i, step in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f"Paso {i}: ")
    run.bold = True
    p.add_run(step)

doc.add_heading("Consejos para categorizar", level=2)
items = [
    "Revisa la descripci\u00f3n del \u00edtem junto con el proveedor (raz\u00f3n social) y su giro",
    "La fecha de emisi\u00f3n ayuda a distinguir gastos de cosecha (nov-dic) vs. trabajos de campo (resto del a\u00f1o)",
    "Si no sabes la subcategor\u00eda exacta, usa \u2018OTRO\u2019 como tipo_gasto",
    "Cada correcci\u00f3n manual mejora el modelo ML para futuras clasificaciones",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# =====================================================================
# 9. IMPORTAR EXCEL
# =====================================================================
doc.add_heading("9. Importar Excel y reentrenar modelo", level=1)
doc.add_paragraph("Ve al men\u00fa Excel > Importar Excel. La aplicaci\u00f3n:")
steps = [
    "Te pide seleccionar el archivo Excel con las correcciones",
    "Compara cada fila con la base de datos original (usa un hash interno para detectar cambios)",
    "Solo actualiza las filas que realmente cambiaron",
    "Marca las filas corregidas con origen=\u2018MANUAL\u2019 y confianza=100%",
    "Valida que las combinaciones categor\u00eda/subcategor\u00eda/tipo_gasto existan en el cat\u00e1logo",
    "Reentrena autom\u00e1ticamente el modelo ML con los datos corregidos",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

p = doc.add_paragraph()
run = p.add_run("Aprendizaje continuo: ")
run.bold = True
run.font.color.rgb = RGBColor(0x0F, 0x66, 0x45)
p.add_run(
    "Cada vez que importas correcciones, el modelo aprende de tus decisiones. "
    "Con el tiempo, la categorizaci\u00f3n autom\u00e1tica ser\u00e1 cada vez m\u00e1s precisa "
    "y necesitar\u00e1s hacer menos correcciones manuales. Con ~200 correcciones "
    "manuales ver\u00e1s una mejora significativa."
)

# =====================================================================
# 10. INVENTARIO
# =====================================================================
doc.add_heading("10. M\u00f3dulo de Inventario", level=1)
doc.add_paragraph(
    "El m\u00f3dulo de inventario permite gestionar el stock de insumos agr\u00edcolas "
    "(fertilizantes, agroqu\u00edmicos, herbicidas, etc.). El stock se construye "
    "autom\u00e1ticamente a partir de las facturas compradas (entradas) y se descuenta "
    "con los usos registrados manualmente."
)

doc.add_heading("10.1 Vista de stock actual", level=2)
doc.add_paragraph(
    "El panel derecho \u2018Stock actual\u2019 muestra una tabla con todos los productos del cat\u00e1logo:"
)
make_table(
    ["Columna", "Descripci\u00f3n"],
    [
        ["C\u00f3digo", "C\u00f3digo \u00fanico del producto (ej.: HERBICIDA-X)"],
        ["Descripci\u00f3n", "Nombre estandarizado del producto"],
        ["Unidad", "Unidad de medida (L, kg, saco, etc.)"],
        ["Stock actual", "Cantidad disponible (entradas \u2212 salidas). En rojo si es negativo."],
        ["Ultima modificaci\u00f3n", "Fecha del \u00faltimo movimiento registrado"],
        ["Tipo", "Tipo del \u00faltimo movimiento (COMPRA, USO_MANUAL, AJUSTE, etc.)"],
    ],
    col_widths=[4, 13],
)
doc.add_paragraph(
    "Usa el campo Buscar para filtrar productos por c\u00f3digo o descripci\u00f3n. "
    "Presiona Enter despu\u00e9s de escribir para aplicar el filtro."
)

doc.add_heading("10.2 Registrar uso manual", level=2)
doc.add_paragraph(
    "Usa el panel izquierdo \u2018Registrar uso\u2019 para descontar del stock:"
)
items = [
    "Selecciona el producto desde el desplegable (c\u00f3digo | descripci\u00f3n [unidad]).",
    "Ingresa la cantidad usada (usa punto o coma como separador decimal).",
    "Confirma o ajusta la fecha (formato YYYY-MM-DD, por defecto hoy).",
    "Agrega una observaci\u00f3n opcional (ej.: \u2018Aplicaci\u00f3n lote norte\u2019).",
    "Presiona Registrar uso.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

add_note(
    "Si el uso deja el stock negativo, la aplicaci\u00f3n mostrar\u00e1 una advertencia "
    "y te preguntar\u00e1 si deseas guardar igualmente."
)

doc.add_heading("10.3 Edici\u00f3n directa de celdas", level=2)
doc.add_paragraph(
    "Haz doble clic sobre cualquier celda editable de la tabla de stock "
    "para modificarla directamente. Columnas editables:"
)
items = [
    "Descripci\u00f3n: renombrar el producto en el cat\u00e1logo.",
    "Unidad: corregir la unidad de medida.",
    "Stock actual: ajuste manual del stock (crea un movimiento tipo AJUSTE_MANUAL).",
    "Ultima modificaci\u00f3n: corregir la fecha del \u00faltimo movimiento.",
    "Tipo: corregir el tipo del \u00faltimo movimiento.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

add_note("La columna C\u00f3digo NO es editable para mantener la integridad referencial con los DTE.")

doc.add_heading("10.4 Sincronizar desde DTE", level=2)
doc.add_paragraph(
    "El bot\u00f3n Refrescar (o men\u00fa Inventario > Refrescar Inventario) re-sincroniza "
    "el cat\u00e1logo y los movimientos de entrada desde los datos de detalle de facturas. "
    "Esta operaci\u00f3n es segura y se puede ejecutar en cualquier momento. "
    "Equivale a la Etapa 6 del pipeline de actualizaci\u00f3n."
)

# =====================================================================
# 11. APLICACIONES DE CAMPO
# =====================================================================
doc.add_heading("11. M\u00f3dulo de Aplicaciones de campo", level=1)
doc.add_paragraph(
    "El m\u00f3dulo de Aplicaciones permite planificar y registrar los tratamientos "
    "fitosanitarios, fertilizaciones y otras labores que impliquen el uso de "
    "productos qu\u00edmicos. Cada aplicaci\u00f3n registrada como EJECUTADA descuenta "
    "autom\u00e1ticamente los productos del inventario."
)

doc.add_heading("11.1 Crear una nueva aplicaci\u00f3n", level=2)
doc.add_paragraph("En el panel izquierdo \u2018Nueva aplicaci\u00f3n de campo\u2019:")
items = [
    "T\u00edtulo: nombre descriptivo de la aplicaci\u00f3n (ej.: \u2018Fungicida lote A\u2019).",
    "Fecha programada: en formato YYYY-MM-DD.",
    "Estado inicial: PROGRAMADA (puede cambiarse a EJECUTADA despu\u00e9s).",
    "Descripci\u00f3n: notas o instrucciones adicionales.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("11.2 Agregar productos qu\u00edmicos", level=2)
doc.add_paragraph("En la secci\u00f3n \u2018Producto qu\u00edmico\u2019 del formulario:")
items = [
    "Selecciona el producto desde el desplegable (muestra solo productos del cat\u00e1logo de inventario).",
    "Ingresa la cantidad a usar.",
    "Agrega una observaci\u00f3n opcional.",
    "Presiona Cargar producto para agregarlo a la lista.",
    "Repite para cada producto necesario en la aplicaci\u00f3n.",
    "Para eliminar un producto de la lista, sel\u00e9ccionalo y presiona Quitar seleccionado.",
    "Cuando est\u00e9s listo, presiona Cargar para guardar la aplicaci\u00f3n.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

add_note(
    "Al guardar la aplicaci\u00f3n como EJECUTADA, los productos se descontar\u00e1n "
    "autom\u00e1ticamente del inventario."
)

doc.add_heading("11.3 Calendario mensual", level=2)
doc.add_paragraph(
    "El panel derecho \u2018Calendario y gesti\u00f3n\u2019 muestra un calendario del mes actual. "
    "Los d\u00edas con aplicaciones programadas aparecen resaltados."
)
items = [
    "Usa los botones < y > para navegar entre meses.",
    "El bot\u00f3n Hoy regresa al mes actual.",
    "El filtro Estado permite ver solo PROGRAMADAS, solo EJECUTADAS o TODOS.",
    "Presiona Refrescar para recargar el calendario desde la base de datos.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

doc.add_heading("11.4 Gestionar aplicaciones existentes", level=2)
doc.add_paragraph(
    "Al hacer clic en un d\u00eda del calendario que tiene aplicaciones, "
    "se mostrar\u00e1n en la lista inferior. Desde ah\u00ed puedes:"
)
items = [
    "Editar una aplicaci\u00f3n existente (carga sus datos en el formulario izquierdo).",
    "Marcar una aplicaci\u00f3n como EJECUTADA.",
    "Eliminar una aplicaci\u00f3n.",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# =====================================================================
# 12. CHEQUEO SEMANAL
# =====================================================================
doc.add_heading("12. Chequeo semanal autom\u00e1tico en segundo plano", level=1)
doc.add_paragraph(
    "Cuando abres Terralix ERP, se inicia autom\u00e1ticamente un proceso en segundo plano "
    "que verifica si hay nuevos DTE en el SII cada cierto n\u00famero de d\u00edas (por defecto 7 d\u00edas). "
    "Si detecta que ha pasado m\u00e1s tiempo del intervalo configurado desde la \u00faltima descarga, "
    "ejecuta el pipeline completo de forma silenciosa."
)
make_table(
    ["Variable en config.env", "Descripci\u00f3n", "Por defecto"],
    [
        ["AUTO_WEEKLY_DTE_CHECK_ENABLED", "Habilita o deshabilita el chequeo autom\u00e1tico", "true"],
        ["AUTO_WEEKLY_DTE_CHECK_INTERVAL_DAYS", "D\u00edas entre chequeos autom\u00e1ticos (1-30)", "7"],
        ["AUTO_WEEKLY_DTE_CHECK_RETRY_HOURS", "Horas antes de reintentar si fall\u00f3 (1-48)", "6"],
        ["AUTO_WEEKLY_DTE_CHECK_POLL_MINUTES", "Frecuencia de verificaci\u00f3n del temporizador (5-1440)", "30"],
    ],
    col_widths=[7, 6.5, 3.5],
)
add_note(
    "El chequeo autom\u00e1tico usa el mismo pipeline de 6 etapas que el bot\u00f3n Actualizar. "
    "Los logs se guardan en %APPDATA%\\Terralix ERP\\logs\\auto_weekly_dte_check.log"
)

# =====================================================================
# 13. CONFIGURAR RUTAS
# =====================================================================
doc.add_heading("13. Configurar rutas", level=1)
doc.add_paragraph(
    "Disponible desde el men\u00fa Opciones > Configurar Rutas (PDF / Base de Datos). "
    "Permite cambiar:"
)
make_table(
    ["Ruta", "Descripci\u00f3n", "Ejemplo"],
    [
        ["Carpeta PDF", "Donde se guardan los PDFs descargados del SII", "C:\\Users\\...\\DTE"],
        ["Base de datos", "Carpeta donde est\u00e1 DteRecibidos_db.db", "C:\\Users\\...\\Base"],
    ],
    col_widths=[3, 8, 6],
)
doc.add_paragraph(
    "Los cambios se guardan inmediatamente en config.env y aplican en la siguiente operaci\u00f3n."
)

# =====================================================================
# 14. POWER BI / EXCEL
# =====================================================================
doc.add_heading("14. Conexi\u00f3n con Power BI / Excel externo", level=1)

doc.add_heading("14.1 Conectar Excel a la base de datos", level=2)
steps = [
    "Abre Excel y ve a Datos > Obtener datos > De otras fuentes > De ODBC",
    "Si no tienes el driver ODBC de SQLite, desc\u00e1rgalo desde sqliteodbc.ch",
    "Alternativa m\u00e1s simple: Datos > Obtener datos > Desde un libro (abre el Excel exportado)",
    "Selecciona las tablas que necesitas (detalle, documentos, cat\u00e1logo, stock, entradas)",
    "Crea una tabla din\u00e1mica con los datos importados",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("14.2 Conectar Power BI", level=2)
steps = [
    "Abre Power BI Desktop",
    "Inicio > Obtener datos > M\u00e1s > Base de datos SQLite",
    "Navega a la ruta de DteRecibidos_db.db",
    "Selecciona las tablas: detalle, documentos, catalogo_costos, insumos_catalogo, insumos_movimientos",
    "Crea las relaciones: detalle.id_doc = documentos.id_doc y detalle.catalogo_costo_id = catalogo_costos.id",
    "Para inventario: insumos_movimientos.codigo = insumos_catalogo.codigo",
    "Ahora puedes crear dashboards con todas las m\u00e9tricas de gastos e inventario",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_heading("14.3 M\u00e9tricas sugeridas para dashboards", level=2)
items = [
    "Gasto total por categor\u00eda y mes",
    "Top 10 proveedores por monto",
    "Distribuci\u00f3n de gastos: cosecha vs. campo vs. administraci\u00f3n",
    "Tendencia mensual de gastos",
    "Detalles sin clasificar pendientes de revisi\u00f3n",
    "Comparativo interanual por categor\u00eda",
    "Stock actual de insumos vs. consumo mensual",
    "Aplicaciones de campo ejecutadas vs. programadas por mes",
]
for item in items:
    doc.add_paragraph(item, style="List Bullet")

# =====================================================================
# 15. ESTRUCTURA DB
# =====================================================================
doc.add_heading("15. Estructura de la base de datos", level=1)
doc.add_paragraph(
    "La base de datos SQLite (DteRecibidos_db.db) contiene las siguientes tablas:"
)

doc.add_heading("15.1 Tabla documentos", level=2)
doc.add_paragraph("Contiene un registro por cada factura o boleta recibida.")
make_table(
    ["Columna", "Tipo", "Descripci\u00f3n"],
    [
        ["id_doc", "TEXT (PK)", "Identificador \u00fanico: TIPO_RUT_FOLIO"],
        ["tipo_doc", "TEXT", "Tipo de documento (33=factura, 34=exenta, etc.)"],
        ["folio", "TEXT", "N\u00famero de folio del documento"],
        ["rut_emisor", "TEXT", "RUT del proveedor"],
        ["razon_social", "TEXT", "Nombre del proveedor"],
        ["giro", "TEXT", "Giro comercial del proveedor"],
        ["fecha_emision", "DATE", "Fecha de emisi\u00f3n del documento"],
        ["monto_total", "REAL", "Monto total del documento"],
        ["IVA", "REAL", "Monto del IVA"],
        ["ruta_pdf", "TEXT", "Ruta al archivo PDF local"],
    ],
    col_widths=[3.5, 2.5, 11],
)

doc.add_heading("15.2 Tabla detalle", level=2)
doc.add_paragraph("Contiene las l\u00edneas de cada documento (\u00edtems facturados).")
make_table(
    ["Columna", "Tipo", "Descripci\u00f3n"],
    [
        ["id_det", "TEXT (PK)", "Identificador \u00fanico de la l\u00ednea"],
        ["id_doc", "TEXT (FK)", "Referencia al documento padre"],
        ["linea", "INTEGER", "N\u00famero de l\u00ednea dentro del documento"],
        ["codigo", "TEXT", "C\u00f3digo del producto (clave para inventario)"],
        ["descripcion", "TEXT", "Descripci\u00f3n del producto/servicio"],
        ["unidad", "TEXT", "Unidad de medida (L, kg, saco, etc.)"],
        ["cantidad", "REAL", "Cantidad"],
        ["precio_unitario", "REAL", "Precio por unidad"],
        ["monto_item", "REAL", "Monto total de la l\u00ednea"],
        ["categoria", "TEXT", "Categor\u00eda de costo asignada"],
        ["subcategoria", "TEXT", "Subcategor\u00eda de costo"],
        ["tipo_gasto", "TEXT", "Tipo espec\u00edfico de gasto"],
        ["catalogo_costo_id", "INTEGER", "ID del cat\u00e1logo de costos"],
        ["confianza_categoria", "INTEGER", "Nivel de confianza (0\u2013100)"],
        ["needs_review", "INTEGER", "1=necesita revisi\u00f3n, 0=ok"],
        ["origen_clasificacion", "TEXT", "Qui\u00e9n clasific\u00f3: REGLA_RS, ML_LOCAL, MANUAL"],
        ["razon_social", "TEXT", "Nombre del proveedor (desnormalizado)"],
        ["giro", "TEXT", "Giro del proveedor (desnormalizado)"],
        ["fecha_emision", "TEXT", "Fecha de emisi\u00f3n (desnormalizado)"],
    ],
    col_widths=[4, 2.5, 10.5],
)

doc.add_heading("15.3 Tabla catalogo_costos", level=2)
doc.add_paragraph("Define las combinaciones v\u00e1lidas de clasificaci\u00f3n contable.")
make_table(
    ["Columna", "Tipo", "Descripci\u00f3n"],
    [
        ["id", "INTEGER (PK)", "Identificador \u00fanico"],
        ["categoria_costo", "TEXT", "Categor\u00eda principal (ej.: COSECHA)"],
        ["subcategoria_costo", "TEXT", "Subcategor\u00eda (ej.: MANO_OBRA_COSECHA)"],
        ["tipo_gasto", "TEXT", "Tipo espec\u00edfico (ej.: MANO_OBRA)"],
    ],
    col_widths=[4, 2.5, 10.5],
)

doc.add_heading("15.4 Tablas de inventario", level=2)
doc.add_paragraph(
    "Las siguientes tablas gestionan el inventario de insumos agr\u00edcolas:"
)
make_table(
    ["Tabla", "Descripci\u00f3n"],
    [
        ["insumos_catalogo", "Cat\u00e1logo maestro de productos: codigo, descripcion_estandar, unidad_base, stock_actual"],
        ["insumos_movimientos", "Hist\u00f3rico de movimientos: codigo, fecha, cantidad, tipo (COMPRA/USO_MANUAL/AJUSTE), observacion"],
        ["aplicaciones_campo", "Registro de aplicaciones: titulo, fecha, estado, descripcion"],
        ["aplicaciones_productos", "Productos qu\u00edmicos de cada aplicaci\u00f3n: id_aplicacion, codigo, cantidad"],
    ],
    col_widths=[5, 12],
)

# =====================================================================
# 16. CATÁLOGO
# =====================================================================
doc.add_heading("16. Cat\u00e1logo de costos", level=1)
doc.add_paragraph(
    "El sistema utiliza 7 categor\u00edas principales para clasificar los gastos:"
)
make_table(
    ["Categor\u00eda", "Descripci\u00f3n", "Ejemplos"],
    [
        ["ADMINISTRACI\u00d3N", "Gastos de oficina, personal, servicios b\u00e1sicos, TI",
         "Contabilidad, sueldos, software, internet"],
        ["COSECHA", "Todo lo relacionado con la cosecha de cerezas",
         "Mano de obra cosecha, bins, flete, packing, abejas"],
        ["ENERG\u00cdA", "Consumo el\u00e9ctrico y servicios de luz",
         "Cuentas CGE, consumo el\u00e9ctrico"],
        ["GASTOS_FINANCIEROS", "Bancos, cr\u00e9ditos, seguros, arriendos financieros",
         "Comisiones bancarias, intereses, seguros"],
        ["INSUMOS_AGR\u00cdCOLAS", "Fertilizantes, agroqu\u00edmicos, herbicidas",
         "Fertilizantes, fungicidas, herbicidas, insecticidas"],
        ["MANTENCI\u00d3N", "Veh\u00edculos, riego, infraestructura, cercos",
         "Combustible, peajes, TAG, filtros riego, reparaciones"],
        ["TRABAJOS_CAMPO", "Labores agr\u00edcolas, poda, servicios, mano de obra",
         "Poda, tractorista, aplicaci\u00f3n herbicida, jornales"],
    ],
    col_widths=[4, 5.5, 7.5],
)

doc.add_heading("Temporalidad de cosecha", level=2)
doc.add_paragraph(
    "Para cerezas en Chile, la cosecha ocurre principalmente entre noviembre y diciembre. "
    "El modelo ML considera la fecha de emisi\u00f3n para distinguir entre gastos de cosecha "
    "y trabajos de campo regulares:"
)
make_table(
    ["Mes", "Temporada", "Efecto en clasificaci\u00f3n"],
    [
        ["Enero \u2013 Febrero", "Postcosecha", "Limpieza, mantenci\u00f3n postcosecha"],
        ["Marzo \u2013 Mayo", "Oto\u00f1o", "Preparaci\u00f3n terreno, poda"],
        ["Junio \u2013 Agosto", "Invierno", "Poda invernal, mantenci\u00f3n riego"],
        ["Septiembre \u2013 Octubre", "Primavera", "Aplicaciones, floraci\u00f3n"],
        ["Noviembre \u2013 Diciembre", "Cosecha", "Mano de obra cosecha, transporte, packing"],
    ],
    col_widths=[4, 3, 10],
)

# =====================================================================
# 17. FAQ
# =====================================================================
doc.add_heading("17. Preguntas frecuentes", level=1)

faqs = [
    ("\u00bfLa app no puede conectarse al SII?",
     "Verifica que las credenciales en config.env (RUT, CLAVE) sean correctas. "
     "Aseg\u00farate de tener conexi\u00f3n a internet estable y que el portal del SII est\u00e9 disponible."),
    ("\u00bfNo puedo iniciar sesi\u00f3n / las credenciales son incorrectas?",
     "Verifica que est\u00e9s usando el email correcto asignado por el administrador. "
     "Si olvidaste la contrase\u00f1a, usa el enlace \u2018\u00bfOlvidaste tu contrase\u00f1a?\u2019 en la pantalla de ingreso."),
    ("\u00bfUn PDF no se lee correctamente?",
     "Ejecuta nuevamente \u2018Actualizar Base de Datos\u2019 para reprocesar los PDFs pendientes. "
     "Si persiste, el PDF puede estar da\u00f1ado o en formato no soportado."),
    ("\u00bfLa categorizaci\u00f3n est\u00e1 equivocada?",
     "Exporta a Excel, corrige la categor\u00eda manualmente y reimporta. "
     "El modelo aprender\u00e1 de tu correcci\u00f3n para el futuro."),
    ("\u00bfQuiero agregar una nueva categor\u00eda?",
     "Agrega la nueva combinaci\u00f3n en la tabla catalogo_costos de la base de datos "
     "usando un editor SQLite (como DB Browser for SQLite)."),
    ("\u00bfD\u00f3nde est\u00e1 la base de datos?",
     "En la ruta configurada en config.env bajo DB_PATH_DTE_RECIBIDOS. "
     "Puedes verla y cambiarla desde Opciones > Configurar Rutas."),
    ("\u00bfPuedo usar la app sin internet?",
     "La descarga del SII, la lectura IA y el inicio de sesi\u00f3n requieren internet. "
     "La categorizaci\u00f3n ML, exportaci\u00f3n Excel, importaci\u00f3n e inventario funcionan sin conexi\u00f3n."),
    ("\u00bfC\u00f3mo hago respaldo de mis datos?",
     "Copia el archivo DteRecibidos_db.db, la carpeta de PDFs y "
     "%APPDATA%\\Terralix ERP\\ a una ubicaci\u00f3n segura. "
     "Incluye data/classifier_dte.pkl para respaldar el modelo entrenado."),
    ("\u00bfEl modelo ML tiene baja confianza?",
     "Esto es normal al inicio. A medida que corrijas m\u00e1s filas y reimportes, "
     "el modelo mejorar\u00e1 progresivamente. Con ~200 correcciones manuales ver\u00e1s una mejora significativa."),
    ("\u00bfEl stock de un insumo est\u00e1 incorrecto?",
     "Haz doble clic en la celda \u2018Stock actual\u2019 de la tabla de inventario para ajustarlo manualmente. "
     "Esto crea un movimiento tipo AJUSTE_MANUAL para mantener el historial."),
    ("\u00bfEl chequeo autom\u00e1tico no funciona?",
     "Verifica que AUTO_WEEKLY_DTE_CHECK_ENABLED=true en config.env. "
     "Revisa el log en %APPDATA%\\Terralix ERP\\logs\\auto_weekly_dte_check.log."),
    ("\u00bfNo veo las pesta\u00f1as Inventario o Aplicaciones?",
     "Aseg\u00farate de estar usando la versi\u00f3n 1.3.0 o superior. "
     "Si la ventana es muy peque\u00f1a, redimensi\u00f3nala para que aparezcan las tres pesta\u00f1as."),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(f"P: {q}")
    run.bold = True
    p2 = doc.add_paragraph()
    p2.add_run(f"R: {a}")
    p2.paragraph_format.space_after = Pt(12)

# =====================================================================
# GUARDAR
# =====================================================================
doc.save(OUT)
print(f"[OK] Manual generado: {OUT}")
